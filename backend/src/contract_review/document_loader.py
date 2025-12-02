"""
文档加载器

支持加载多种格式的文档：
- 文本格式：.md, .txt
- Word 文档：.docx
- Excel 文档：.xlsx
- PDF 文档：.pdf（智能判断：先尝试提取文本，若文本少则走 OCR）
- 图片格式：.jpg, .jpeg, .png, .webp（OCR）
- 表格格式：.xlsx, .xls, .csv（用于审核标准）

只有图片文件和扫描版 PDF 才会真正调用 OCR API，其他都是本地文本提取。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Optional, Sequence, TYPE_CHECKING

import pandas as pd
import pdfplumber
from docx import Document

from .models import LoadedDocument

if TYPE_CHECKING:
    from .ocr_service import OCRService

logger = logging.getLogger(__name__)

# 支持的文档格式（待审阅文档）
DOCUMENT_EXTENSIONS = {".md", ".txt", ".docx", ".xlsx", ".pdf"}
# 支持的图片格式
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
# 支持的表格格式（用于审核标准）
TABLE_EXTENSIONS = {".xlsx", ".xls", ".csv"}
# 所有支持的格式
SUPPORTED_EXTENSIONS = DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS | TABLE_EXTENSIONS


def scan_documents(root: Path, extensions: Optional[set] = None) -> List[Path]:
    """
    递归扫描目录，返回支持的文件列表

    Args:
        root: 扫描根目录
        extensions: 允许的扩展名集合，None 时使用默认值

    Returns:
        文件路径列表（按名称排序）
    """
    if extensions is None:
        extensions = DOCUMENT_EXTENSIONS

    files: List[Path] = []
    for path in root.rglob("*"):
        if path.is_file() and path.suffix.lower() in extensions:
            files.append(path)
    return sorted(files)


def load_documents(paths: Sequence[Path]) -> List[LoadedDocument]:
    """批量加载文档"""
    return [load_document(path) for path in paths]


def load_document(path: Path) -> LoadedDocument:
    """
    加载单个文档（同步版本，不支持图片和扫描 PDF 的 OCR）

    Args:
        path: 文档路径

    Returns:
        LoadedDocument 对象

    Raises:
        ValueError: 不支持的文件类型
    """
    suffix = path.suffix.lower()

    if suffix in {".md", ".txt"}:
        text = _read_text_file(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".xlsx":
        text = _read_xlsx(path)
    else:
        raise ValueError(f"不支持的文件类型: {path}")

    return LoadedDocument(
        path=path,
        text=text.strip(),
        metadata={
            "filename": path.name,
            "relative_path": str(path),
            "suffix": suffix,
            "size_bytes": path.stat().st_size,
        },
    )


async def load_document_async(
    path: Path,
    ocr_service: Optional["OCRService"] = None
) -> LoadedDocument:
    """
    异步加载单个文档，支持 OCR

    对于图片和扫描版 PDF，需要提供 ocr_service 参数。

    处理策略：
    - .md/.txt: 直接读取文本
    - .docx: 用 python-docx 提取文本
    - .xlsx: 用 openpyxl 提取并转 Markdown 表格
    - .pdf: 智能判断，先尝试提取文本，若文本 < 100 字符则走 OCR
    - 图片: 必须走 OCR

    Args:
        path: 文档路径
        ocr_service: OCR 服务实例（处理图片和扫描 PDF 时需要）

    Returns:
        LoadedDocument 对象

    Raises:
        ValueError: 不支持的文件类型或需要 OCR 但未提供服务
    """
    suffix = path.suffix.lower()

    # 纯文本文件直接读取
    if suffix in {".md", ".txt"}:
        text = _read_text_file(path)

    # Word 文档本地提取
    elif suffix == ".docx":
        text = _read_docx(path)

    # Excel 文档本地提取
    elif suffix == ".xlsx":
        text = _read_xlsx(path)

    # PDF: 智能判断
    elif suffix == ".pdf":
        if ocr_service:
            # 使用 OCR 服务（内部会智能判断是否需要 OCR）
            text = await ocr_service.ocr_pdf(str(path))
        else:
            # 没有 OCR 服务，使用传统方式
            text = _read_pdf(path)

    # 图片: 必须 OCR
    elif suffix in IMAGE_EXTENSIONS:
        if not ocr_service:
            raise ValueError(f"处理图片文件需要配置 OCR 服务: {path}")
        text = await ocr_service.ocr_image(str(path))

    else:
        raise ValueError(f"不支持的文件类型: {path}")

    return LoadedDocument(
        path=path,
        text=text.strip(),
        metadata={
            "filename": path.name,
            "relative_path": str(path),
            "suffix": suffix,
            "size_bytes": path.stat().st_size,
            "ocr_used": suffix in IMAGE_EXTENSIONS or (suffix == ".pdf" and ocr_service is not None),
        },
    )


def _read_text_file(path: Path) -> str:
    """读取纯文本文件"""
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    """读取 Word 文档"""
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]

    # 也读取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    """读取 PDF 文档"""
    contents: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            contents.append(page_text)
    return "\n".join(contents)


def _read_xlsx(path: Path) -> str:
    """读取 Excel 文档并转换为 Markdown 表格"""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("请安装 openpyxl: pip install openpyxl")

    wb = openpyxl.load_workbook(path, data_only=True)
    results = []

    for sheet in wb.worksheets:
        sheet_content = [f"## Sheet: {sheet.title}"]

        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # 构建 Markdown 表格
        for i, row in enumerate(rows):
            cells = [str(c) if c is not None else "" for c in row]
            sheet_content.append("| " + " | ".join(cells) + " |")
            if i == 0:
                # 添加表头分隔线
                sheet_content.append("| " + " | ".join(["---"] * len(cells)) + " |")

        results.append("\n".join(sheet_content))

    return "\n\n".join(results)


# ==================== 表格读取函数 ====================

def load_table(path: Path) -> pd.DataFrame:
    """
    加载表格文件

    Args:
        path: 表格文件路径（.xlsx, .xls, .csv）

    Returns:
        pandas DataFrame

    Raises:
        ValueError: 不支持的文件类型
    """
    suffix = path.suffix.lower()

    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(path)
    elif suffix == ".csv":
        # 尝试不同编码
        for encoding in ["utf-8", "gbk", "gb2312", "utf-8-sig"]:
            try:
                return pd.read_csv(path, encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法解析 CSV 文件编码: {path}")
    else:
        raise ValueError(f"不支持的表格格式: {path}")


def is_document_file(path: Path) -> bool:
    """检查是否为文档文件"""
    return path.suffix.lower() in DOCUMENT_EXTENSIONS


def is_image_file(path: Path) -> bool:
    """检查是否为图片文件"""
    return path.suffix.lower() in IMAGE_EXTENSIONS


def is_table_file(path: Path) -> bool:
    """检查是否为表格文件"""
    return path.suffix.lower() in TABLE_EXTENSIONS


def get_file_type(path: Path) -> Optional[str]:
    """
    获取文件类型

    Returns:
        "document" | "image" | "table" | None
    """
    suffix = path.suffix.lower()
    if suffix in DOCUMENT_EXTENSIONS:
        return "document"
    elif suffix in IMAGE_EXTENSIONS:
        return "image"
    elif suffix in TABLE_EXTENSIONS:
        return "table"
    return None
