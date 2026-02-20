from pathlib import Path

import pytest

from contract_review.document_loader import load_document
from contract_review.plugins.fidic import register_fidic_plugin
from contract_review.plugins.registry import clear_plugins, get_parser_config
from contract_review.structure_parser import StructureParser


class TestUploadPipeline:
    def test_load_and_parse_txt(self):
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
            f.write("1.1 Definitions\nThe Employer means the party...\n")
            f.write("1.2 Interpretation\nWords importing...\n")
            f.write("4.1 Contractor Obligations\nThe Contractor shall...\n")
            tmp_path = f.name

        loaded = load_document(Path(tmp_path))
        assert len(loaded.text) > 0

        parser = StructureParser()
        structure = parser.parse(loaded)
        assert structure.total_clauses >= 1

        Path(tmp_path).unlink(missing_ok=True)

    def test_load_and_parse_docx(self):
        pytest.importorskip("docx")

        from docx import Document
        import tempfile

        doc = Document()
        doc.add_paragraph("1.1 Definitions")
        doc.add_paragraph("The Employer means the party named in the Contract.")
        doc.add_paragraph("1.2 Obligations")
        doc.add_paragraph("The Contractor shall perform the Works.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        loaded = load_document(Path(tmp_path))
        assert len(loaded.text) > 0

        parser = StructureParser()
        structure = parser.parse(loaded)
        assert structure.total_clauses >= 1

        Path(tmp_path).unlink(missing_ok=True)

    def test_parser_config_from_plugin(self):
        clear_plugins()
        register_fidic_plugin()

        config = get_parser_config("fidic")
        assert config.structure_type == "fidic_gc"
        assert config.definitions_section_id == "1.1"
